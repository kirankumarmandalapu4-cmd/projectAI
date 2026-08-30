import os
import shutil
import subprocess
from pathlib import Path
from typing import List, Dict, Any

class DocumentLoader:
    TEXT_EXTENSIONS = {
        "txt", "csv", "md", "markdown", "json", "xml", "html", "htm", "rtf", "log"
    }
    IMAGE_EXTENSIONS = {"png", "jpg", "jpeg", "webp", "bmp", "tif", "tiff", "gif"}

    @staticmethod
    def load_file(file_path: str, file_type: str) -> List[Dict[str, Any]]:
        """
        Extract text content page by page from supported office, text, PDF, and image files.
        Returns list of dicts: [{'page_number': int, 'text': str}]
        """
        ext = file_type.lower().lstrip(".")
        if ext == "pdf":
            return DocumentLoader._load_pdf(file_path)
        elif ext == "docx":
            return DocumentLoader._load_docx(file_path)
        elif ext == "doc":
            return DocumentLoader._load_legacy_doc(file_path)
        elif ext in DocumentLoader.IMAGE_EXTENSIONS:
            return DocumentLoader._load_image(file_path)
        elif ext in DocumentLoader.TEXT_EXTENSIONS:
            return DocumentLoader._load_txt(file_path)
        else:
            raise ValueError(f"Unsupported file format: {file_type}")

    @staticmethod
    def _load_pdf(file_path: str) -> List[Dict[str, Any]]:
        pages = []
        try:
            import fitz # PyMuPDF
            doc = fitz.open(file_path)
            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                text = page.get_text("text")
                if not text or not text.strip():
                    # OCR is optional and only runs for scanned/ image-only
                    # pages. A normal text PDF never pays this cost.
                    try:
                        import pytesseract
                        from PIL import Image
                        pixmap = page.get_pixmap(matrix=fitz.Matrix(1.5, 1.5), alpha=False)
                        image = Image.frombytes("RGB", [pixmap.width, pixmap.height], pixmap.samples)
                        text = pytesseract.image_to_string(image)
                    except Exception:
                        text = ""
                if text and text.strip():
                    pages.append({
                        "page_number": page_num + 1,
                        "text": text.strip()
                    })
            doc.close()
        except ImportError as e:
            raise RuntimeError("PDF support is unavailable. Install the pymupdf package.") from e
        except Exception as e:
            raise ValueError(f"Could not read PDF file: {e}") from e
        return pages

    @staticmethod
    def _load_docx(file_path: str) -> List[Dict[str, Any]]:
        pages = []
        try:
            import docx
            doc = docx.Document(file_path)
            parts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
            for table in doc.tables:
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    if any(cells):
                        parts.append(" | ".join(cells))
            full_text = "\n".join(parts)
            pages.append({"page_number": 1, "text": full_text})
        except ImportError as e:
            raise RuntimeError("Word support is unavailable. Install the python-docx package.") from e
        except Exception as e:
            raise ValueError(f"Could not read Word document: {e}") from e
        return pages

    @staticmethod
    def _load_legacy_doc(file_path: str) -> List[Dict[str, Any]]:
        """Read legacy .doc files when a local converter is installed."""
        converter = shutil.which("antiword") or shutil.which("soffice") or shutil.which("libreoffice")
        if not converter:
            raise ValueError(
                "Legacy .doc files require LibreOffice or antiword. Save the file as .docx and try again."
            )

        try:
            if Path(converter).stem.lower() in {"soffice", "libreoffice"}:
                result = subprocess.run(
                    [converter, "--headless", "--convert-to", "txt:Text", "--outdir", str(Path(file_path).parent), file_path],
                    capture_output=True, text=True, timeout=60, check=False,
                )
                converted_path = Path(file_path).with_suffix(".txt")
                if result.returncode != 0 or not converted_path.exists():
                    raise ValueError(result.stderr.strip() or "LibreOffice could not convert the .doc file.")
                return DocumentLoader._load_txt(str(converted_path))

            result = subprocess.run([converter, file_path], capture_output=True, text=True, timeout=60, check=False)
            if result.returncode != 0:
                raise ValueError(result.stderr.strip() or "antiword could not read the .doc file.")
            return [{"page_number": 1, "text": result.stdout}]
        except subprocess.TimeoutExpired as e:
            raise ValueError("Reading the .doc file timed out.") from e

    @staticmethod
    def _load_image(file_path: str) -> List[Dict[str, Any]]:
        """Extract text from an uploaded image using optional local OCR."""
        try:
            from PIL import Image
            with Image.open(file_path) as image:
                image.load()
                width, height = image.size
                image_format = image.format or Path(file_path).suffix.lstrip(".").upper()
                try:
                    import pytesseract
                    text = pytesseract.image_to_string(image).strip()
                except Exception:
                    text = ""

                if not text:
                    text = (
                        f"Uploaded image: {Path(file_path).name}\n"
                        f"Image type: {image_format}\n"
                        f"Image dimensions: {width} x {height} pixels.\n"
                        "No readable OCR text was detected. Install Tesseract OCR for image text extraction."
                    )
                return [{"page_number": 1, "text": text}]
        except ImportError as e:
            raise RuntimeError("Image support is unavailable. Install Pillow and pytesseract.") from e
        except Exception as e:
            raise ValueError(f"Could not read image file: {e}") from e

    @staticmethod
    def _load_txt(file_path: str) -> List[Dict[str, Any]]:
        pages = []
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            content = f.read()
        pages.append({"page_number": 1, "text": content})
        return pages
